from flask import Flask, render_template, flash, request, session
from flask import render_template, redirect, url_for, request
#from wtforms import Form, TextField, TextAreaField, validators, StringField, SubmitField
#from werkzeug.utils import secure_filename
from PIL import Image, ImageChops,ImageStat
import yagmail
from tkinter import *
import cv2
import csv
import os
import time
import numpy as np
from PIL import Image,ImageTk

import mysql.connector

import datetime
app = Flask(__name__)
app.config['DEBUG']
app.config['SECRET_KEY'] = '7d441f27d441f27567d441f2b6176a'

def dmail():
    ts = time.time()
    date = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d')
    import mysql.connector
    conn = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')
    cursor = conn.cursor()
    cur = conn.cursor()
    cur.execute("SELECT id,Userid,Status FROM attentb where date='"+date+"' and Status='absent'")
    data = cur.fetchall()
    print(data)

    # Import docx NOT python-docx
    import docx

    # Create an instance of a word document
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading('Student Attendance', 0)

    # Table data in a form of list
    # data = ((1, 'Geek 1'),(2, 'Geek 2'),(3, 'Geek 3'))

    # Creating a table object
    table = doc.add_table(rows=1, cols=3)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'RegNo'
    row[2].text = 'Status'

    # Adding data from the list to the table
    for id, RegNo, Status in data:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = RegNo
        row[2].text = Status

    # Now save the document to a location
    table.style = 'Colorful List'
    doc.save('record.docx')

@app.route("/")
def homepage():
    return render_template('index.html')
@app.route("/Home")
def Home():
    return render_template('index.html')


@app.route("/AdminLogin")
def AdminLogin():
    return render_template('AdminLogin.html')


@app.route("/atmport")
def finger():
    return render_template('atmport.html')




@app.route("/adminlogin", methods=['GET', 'POST'])
def adminlogin():
    error = None
    if request.method == 'POST':
       if request.form['uname'] == 'admin' or request.form['password'] == 'admin':
           conn = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')
           cursor = conn.cursor()
           cur = conn.cursor()
           cur.execute("SELECT * FROM register")
           data = cur.fetchall()
           return render_template('AdminHome.html', data=data)

       else:
        return render_template('index.html', error=error)

@app.route("/AdminHome")
def AdminHome():
    conn = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')


    cur = conn.cursor()
    cur.execute("SELECT * FROM register")
    data = cur.fetchall()
    return render_template('AdminHome.html', data=data)







@app.route("/NewStudent")
def NewStudent():

    return render_template('NewUser.html')


@app.route("/NewStudent1", methods=['GET', 'POST'])
def NewStudent1():
     if request.method == 'POST':
         fname = request.form['fname']
         uid = request.form['uid']
         lname = request.form['lname']
         age = request.form['age']
         address = request.form['address']
         pnumber = request.form['pnumber']
         email = request.form['email']
         country = request.form['country']
         bname = request.form['bname']
         acno = request.form['acno']
         ano = request.form['ano']
         pno = request.form['pno']
         state = request.form['state']
         city = request.form['city']
         # initializing string

         import math, random

         # function to generate OTP
         def generateaccno():

             # Declare a digits variable
             # which stores all digits
             digits = "0123456789"
             OTP = ""

             # length of password can be changed
             # by changing value in range
             for i in range(14):
                 OTP += digits[math.floor(random.random() * 10)]

             return OTP
         def generateatmno():

             # Declare a digits variable
             # which stores all digits
             digits = "0123456789"
             OTP = ""

             # length of password can be changed
             # by changing value in range
             for i in range(16):
                 OTP += digits[math.floor(random.random() * 10)]

             return OTP
         def generateatmpin():

             # Declare a digits variable
             # which stores all digits
             digits = "0123456789"
             OTP = ""

             # length of password can be changed
             # by changing value in range
             for i in range(4):
                 OTP += digits[math.floor(random.random() * 10)]

             return OTP

        # acno = generateaccno()
         atmno = generateatmno()
         atmpin = generateatmpin()

         conn = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')
         cursor = conn.cursor()
         cursor.execute(
             "INSERT INTO register VALUES ('','"+uid+"','" + fname + "','" + lname + "','" + age + "','" + address + "','" + pnumber + "','" + email + "','" + country + "','" + state + "','" + city + "','" + bname + "','"+str(acno)+"','" + ano + "','" + pno + "','','"+str(atmno)+"','" + str(atmpin) + "','')")
         conn.commit()
         conn.close()
         cam = cv2.VideoCapture(0)
         cam.set(3, 640)  # set video width
         cam.set(4, 480)  # set video height

         # make sure 'haarcascade_frontalface_default.xml' is in the same folder as this code
         face_detector = cv2.CascadeClassifier('haarcascade_frontalface_default.xml')

         # For each person, enter one numeric face id (must enter number start from 1, this is the lable of person 1)

         face_id = uid
         print(face_id)

         print("\n [INFO] Initializing face capture. Look the camera and wait ...")
         # Initialize individual sampling face count
         count = 0
         while (True):

             ret, img = cam.read()
             gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
             faces = face_detector.detectMultiScale(gray, 1.3, 5)

             for (x, y, w, h) in faces:
                 cv2.rectangle(img, (x, y), (x + w, y + h), (255, 0, 0), 2)
                 count += 1

                 # Save the captured image into the datasets folder
                 cv2.imwrite("dataset/User." + str(face_id) + '.' + str(count) + ".jpg", gray[y:y + h, x:x + w])

                 cv2.imshow('image', img)

             k = cv2.waitKey(100) & 0xff  # Press 'ESC' for exiting video
             if k == 27:
                 break
             elif count >= 30:  # Take 30 face sample and stop video
                 break
         # Do a bit of cleanup
         print("\n [INFO] Exiting Program and cleanup stuff")
         cam.release()
         cv2.destroyAllWindows()
         path = 'dataset'

         recognizer = cv2.face.LBPHFaceRecognizer_create()
         detector = cv2.CascadeClassifier("haarcascade_frontalface_default.xml");

         def getImagesAndLabels(path):

             imagePaths = [os.path.join(path, f) for f in os.listdir(path)]
             faceSamples = []
             ids = []

             for imagePath in imagePaths:

                 PIL_img = Image.open(imagePath).convert('L')  # convert it to grayscale
                 img_numpy = np.array(PIL_img, 'uint8')

                 id = int(os.path.split(imagePath)[-1].split(".")[1])
                 faces = detector.detectMultiScale(img_numpy)

                 for (x, y, w, h) in faces:
                     faceSamples.append(img_numpy[y:y + h, x:x + w])
                     ids.append(id)

             return faceSamples, ids

         print("\n [INFO] Training faces. It will take a few seconds. Wait ...")
         faces, ids = getImagesAndLabels(path)
         recognizer.train(faces, np.array(ids))

         # Save the model into trainer/trainer.yml
         recognizer.write('trainer/trainer.yml')  # recognizer.save() worked on Mac, but not on Pi

         # Print the numer of faces trained and end program
         print("\n [INFO] {0} faces trained. Exiting Program".format(len(np.unique(ids))))

         mail = 'testsam360@gmail.com';
         password = 'rddwmbynfcbgpywf';
         # list of email_id to send the mail
         conn2 = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')


         dest = email
         body = "ATM Card Details Your Crad.No:"+str(atmno)+",and Login Pin:"+str(atmpin)
         yag = yagmail.SMTP(mail, password)
         yag.send(to=dest, subject="Generate ATM Cardnumber and Pin Details ...!", contents=body)

         #dmail()
         flash("Logged in successfully.")

         return 'Register successfully'

         return render_template('AdminHome.html', data=data)


@app.route("/atmverify", methods=['GET', 'POST'])
def atmverify():
     if request.method == 'POST':
         atmno = request.form['atmno']
         session['atmid']=atmno
         conn = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')
         cursor = conn.cursor()
         cursor.execute("select * from register where atmno='"+str(atmno)+"'")
         data = cursor.fetchone()
         if data is None:
             return "No Record"
         else:
             now = time.time()  ###For calculate seconds of video
             future = now + 120

             my_list = ['']

             conn = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')
             cursor = conn.cursor()
             cursor.execute("select * from register")
             data = cursor.fetchall()
             print(data)
             for data1 in data:
                 my_list.append(data1[2])
             print(my_list)

             import cv2
             import numpy as np
             import os

             recognizer = cv2.face.LBPHFaceRecognizer_create()
             recognizer.read('trainer/trainer.yml')  # load trained model
             cascadePath = "haarcascade_frontalface_default.xml"
             faceCascade = cv2.CascadeClassifier(cascadePath);

             font = cv2.FONT_HERSHEY_SIMPLEX

             # iniciate id counter, the number of persons you want to include
             id = 0  # two persons (e.g. Jacob, Jack)

             names = my_list  # key in names, start from the second place, leave first empty

             # Initialize and start realtime video capture
             cam = cv2.VideoCapture(0)
             cam.set(3, 640)  # set video widht
             cam.set(4, 480)  # set video height

             # Define min window size to be recognized as a face
             minW = 0.1 * cam.get(3)
             minH = 0.1 * cam.get(4)

             while True:

                 ret, img = cam.read()
                 cv2.imwrite("data.png", img)
                 cv2.imwrite("C:/wamp/wwww/login-form-14/images/bg-1.jpg", img)


                 gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

                 faces = faceCascade.detectMultiScale(
                     gray,
                     scaleFactor=1.2,
                     minNeighbors=5,
                     minSize=(int(minW), int(minH)),
                 )

                 for (x, y, w, h) in faces:

                     cv2.rectangle(img, (x, y), (x + w, y + h), (0, 255, 0), 2)

                     id, confidence = recognizer.predict(gray[y:y + h, x:x + w])

                     # Check if confidence is less them 100 ==> "0" is perfect match
                     if (confidence < 50):
                         id1 = id

                         # id = names[id]
                         confidence = "  {0}%".format(round(100 - confidence))


                     else:
                         id = "unknown"
                         confidence = "  {0}%".format(round(100 - confidence))

                     cv2.putText(img, str(id), (x + 5, y - 5), font, 1, (255, 255, 255), 2)
                     cv2.putText(img, str(confidence), (x + 5, y + h - 5), font, 1, (255, 255, 0), 1)

                 cv2.imshow('camera', img)
                 print(id)

                 if time.time() > future:
                     break

                 k = cv2.waitKey(10) & 0xff  # Press 'ESC' for exiting video
                 if k == 27:
                     break

             # Do a bit of cleanup
             print("\n [INFO] Exiting Program and cleanup stuff")
             cam.release()
             cv2.destroyAllWindows()
             cursor1 = conn.cursor()
             cursor1.execute("select * from register where uid='"+str(id)+"' and atmno='"+str(atmno)+"'")
             data2 = cursor1.fetchone()
             if data2 is None:
                 cursor1 = conn.cursor()
                 cursor1.execute("select * from register where atmno='" + str(atmno) + "'")
                 data2 = cursor1.fetchone()
                 i=data2[0]

                 email=data2[7]
                 mail = 'testsam360@gmail.com';
                 password = 'rddwmbynfcbgpywf';
                 # list of email_id to send the mail
                 li = [email]
                 body = "http://192.168.43.205/login-form-14/index.php?id="+str(i)
                 yag = yagmail.SMTP(mail, password)
                 yag.send(to=email,subject="Alert...!",contents=body, attachments="data.png")


                 return render_template("status.html",data='unknown',status='Waiting')

             else:
                 return render_template("atmpin.html")

@app.route("/atmpin", methods=['GET', 'POST'])
def atmpin():
    atmpin=request.form['atmno']
    atmno=session['atmid']

    conn = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')
    cursor = conn.cursor()
    cursor.execute("select * from register where atmno='" + str(atmno) + "' and atmpin='"+str(atmpin)+"'")
    data = cursor.fetchone()
    if data is None:
        return "No Record"
    else:
        return render_template("amount.html")
@app.route("/amount", methods=['GET', 'POST'])
def amount():
    amount=request.form['amount']



    return render_template("result.html",data=amount)
@app.route("/Useraccess", methods=['GET', 'POST'])
def Useraccess():
    amount=request.form['clicked_btn']
    print(amount)



    return render_template("result.html",data=amount)
@app.route("/status")
def status():
    i=session['atmid']
    #print(amount)
    conn = mysql.connector.connect(user='root', password='', host='localhost', database='Bankatm')
    cursor = conn.cursor()
    cursor.execute("select * from register where atmno='" + str(i) + "'")
    data = cursor.fetchone()
    amount=data[18]
    status=data[15]
    if status=='Accepted':
        status='User Accepted Your Request And Amount Transfer Success'
    else:
        status=status



    return render_template("status.html",amount=amount,status=status)


if __name__ == '__main__':
    app.run(debug=True, use_reloader=True)